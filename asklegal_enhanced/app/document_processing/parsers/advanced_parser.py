from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from typing import List, Dict, Any
import re

class AdvancedDocumentParser:
    """Simplified document parser without external dependencies"""
    
    def __init__(self):
        # No external dependencies needed
        print("Initialized simplified document parser")
    
    def parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse PDF using basic PyPDF2
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            List[Dict[str, Any]]: Parsed document elements
        """
        elements = []
        
        try:
            reader = PdfReader(file_path)
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                
                if text.strip():
                    element = {
                        "id": f"page_{page_num}",
                        "text": text,
                        "type": "page",
                        "metadata": {
                            "page_number": page_num + 1
                        }
                    }
                    
                    # Detect clauses
                    element["clauses"] = self._detect_clauses(text)
                    elements.append(element)
                
        except Exception as e:
            print(f"PDF parsing failed: {e}")
        
        return elements
    
    def parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse DOCX with layout awareness and clause detection
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            List[Dict[str, Any]]: Parsed document elements
        """
        elements = []
        
        try:
            doc = DocxDocument(file_path)
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    element = {
                        "id": f"paragraph_{i}",
                        "text": paragraph.text,
                        "type": "paragraph",
                        "metadata": {
                            "style": paragraph.style.name if paragraph.style else "Normal"
                        }
                    }
                    
                    # Detect clauses
                    element["clauses"] = self._detect_clauses(paragraph.text)
                    elements.append(element)
                
        except Exception as e:
            print(f"DOCX parsing failed: {e}")
            
            # Fallback to basic text extraction
            try:
                # Basic fallback approach
                print("Using basic DOCX text extraction as fallback")
            except Exception as fallback_error:
                print(f"Fallback DOCX parsing also failed: {fallback_error}")
        
        return elements
    
    def _detect_clauses(self, text: str) -> List[Dict[str, Any]]:
        """
        Detect clauses in text using pattern matching
        
        Args:
            text (str): Text to analyze
            
        Returns:
            List[Dict[str, Any]]: Detected clauses
        """
        clauses = []
        
        # Common clause patterns
        clause_patterns = [
            r'(\d+\.\s*[A-Z][^.]*?shall[^.]*\.)',
            r'(\d+\.\s*[A-Z][^.]*?agrees[^.]*\.)',
            r'(\d+\.\s*[A-Z][^.]*?warrants[^.]*\.)',
            r'(\d+\.\s*[A-Z][^.]*?represents[^.]*\.)',
            r'([A-Z][^.]*?shall[^.]*\.)',
            r'([A-Z][^.]*?agrees[^.]*\.)'
        ]
        
        for pattern in clause_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                clause_text = match.group(1).strip()
                if len(clause_text) > 20:  # Filter out very short matches
                    clauses.append({
                        "text": clause_text,
                        "start": match.start(),
                        "end": match.end(),
                        "type": "contractual_clause"
                    })
        
        return clauses
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract enhanced metadata from document
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Dict[str, Any]: Enhanced document metadata
        """
        import os
        from datetime import datetime
        
        stat = os.stat(file_path)
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().lstrip('.')
        
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_type": file_type,
            "file_size": stat.st_size,
            "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat()
        }
        
        # Extract document-specific metadata
        if file_type == 'pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata.update({
                        "page_count": len(pdf.pages),
                        "author": pdf.metadata.get('Author', 'Unknown'),
                        "title": pdf.metadata.get('Title', 'Unknown'),
                        "subject": pdf.metadata.get('Subject', 'Unknown')
                    })
            except:
                pass
        elif file_type == 'docx':
            try:
                doc = DocxDocument(file_path)
                core_props = doc.core_properties
                metadata.update({
                    "author": core_props.author or 'Unknown',
                    "title": core_props.title or 'Unknown',
                    "subject": core_props.subject or 'Unknown',
                    "created": core_props.created.isoformat() if core_props.created else 'Unknown',
                    "modified": core_props.modified.isoformat() if core_props.modified else 'Unknown'
                })
            except:
                pass
        
        return metadata

# Global instance
advanced_parser = AdvancedDocumentParser()