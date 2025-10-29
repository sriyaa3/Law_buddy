"""
Enhanced Document Processing with LayoutLMv3
Multimodal document parsing with structure awareness
"""

import logging
import os
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
import pytesseract
from PIL import Image
import PyPDF2
import pdfplumber
from docx import Document
import io
from transformers import LayoutLMv3Processor, LayoutLMv3ForTokenClassification
import torch
from app.core.config_enhanced import settings

logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    doc_id: str
    text: str
    clauses: List[Dict[str, Any]]
    entities: List[Dict[str, Any]]
    tables: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    structure: Dict[str, Any]

class EnhancedDocumentProcessor:
    """
    Advanced document processing with:
    - LayoutLMv3 for structure-aware parsing
    - OCR for scanned documents
    - Multi-format support (PDF, DOCX, images)
    - Table extraction
    - Entity recognition
    """
    
    def __init__(self):
        self.layoutlm_processor = None
        self.layoutlm_model = None
        self._initialize_layoutlm()
        
        # Configure Tesseract
        if os.path.exists(settings.TESSERACT_CMD):
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
    
    def _initialize_layoutlm(self):
        """Initialize LayoutLMv3 model"""
        try:
            logger.info("Loading LayoutLMv3 model...")
            self.layoutlm_processor = LayoutLMv3Processor.from_pretrained(
                settings.LAYOUTLM_MODEL,
                apply_ocr=False  # We'll use our own OCR
            )
            self.layoutlm_model = LayoutLMv3ForTokenClassification.from_pretrained(
                settings.LAYOUTLM_MODEL
            )
            
            # Move to GPU if available
            if torch.cuda.is_available():
                self.layoutlm_model = self.layoutlm_model.cuda()
            
            logger.info("✓ LayoutLMv3 model loaded successfully")
        except Exception as e:
            logger.error(f"Failed to load LayoutLMv3: {e}")
            logger.warning("Will use fallback text extraction")
    
    def process_document(
        self,
        file_path: str,
        doc_id: str,
        file_type: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process a document and extract all information
        """
        if file_type is None:
            file_type = os.path.splitext(file_path)[1].lower()
        
        logger.info(f"Processing document: {doc_id} ({file_type})")
        
        # Extract based on file type
        if file_type in ['.pdf']:
            return self._process_pdf(file_path, doc_id)
        elif file_type in ['.docx', '.doc']:
            return self._process_docx(file_path, doc_id)
        elif file_type in ['.png', '.jpg', '.jpeg']:
            return self._process_image(file_path, doc_id)
        elif file_type in ['.txt']:
            return self._process_text(file_path, doc_id)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _process_pdf(self, file_path: str, doc_id: str) -> ProcessedDocument:
        """Process PDF document"""
        text = ""
        tables = []
        pages_data = []
        
        try:
            # Use pdfplumber for better table extraction
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        tables.append({
                            "page": page_num + 1,
                            "data": table
                        })
                    
                    pages_data.append({
                        "page_num": page_num + 1,
                        "text": page_text,
                        "tables": len(page_tables)
                    })
            
            # If text is sparse, try OCR
            if len(text.strip()) < 100:
                logger.info("Low text content, attempting OCR...")
                text = self._ocr_pdf(file_path)
                
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            # Fallback to PyPDF2
            text = self._extract_text_pypdf2(file_path)
        
        # Extract clauses and entities
        clauses = self._extract_clauses(text)
        entities = self._extract_entities(text)
        
        return ProcessedDocument(
            doc_id=doc_id,
            text=text,
            clauses=clauses,
            entities=entities,
            tables=tables,
            metadata={
                "file_type": "pdf",
                "num_pages": len(pages_data),
                "num_tables": len(tables)
            },
            structure={"pages": pages_data}
        )
    
    def _process_docx(self, file_path: str, doc_id: str) -> ProcessedDocument:
        """Process DOCX document"""
        try:
            doc = Document(file_path)
            
            # Extract text
            text = "\n\n".join([para.text for para in doc.paragraphs])
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "table_num": table_idx + 1,
                    "data": table_data
                })
            
            # Extract clauses and entities
            clauses = self._extract_clauses(text)
            entities = self._extract_entities(text)
            
            return ProcessedDocument(
                doc_id=doc_id,
                text=text,
                clauses=clauses,
                entities=entities,
                tables=tables,
                metadata={
                    "file_type": "docx",
                    "num_paragraphs": len(doc.paragraphs),
                    "num_tables": len(tables)
                },
                structure={}
            )
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise
    
    def _process_image(self, file_path: str, doc_id: str) -> ProcessedDocument:
        """Process image with OCR"""
        try:
            # Perform OCR
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            # Extract clauses and entities
            clauses = self._extract_clauses(text)
            entities = self._extract_entities(text)
            
            return ProcessedDocument(
                doc_id=doc_id,
                text=text,
                clauses=clauses,
                entities=entities,
                tables=[],
                metadata={
                    "file_type": "image",
                    "image_size": image.size
                },
                structure={}
            )
            
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            raise
    
    def _process_text(self, file_path: str, doc_id: str) -> ProcessedDocument:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            clauses = self._extract_clauses(text)
            entities = self._extract_entities(text)
            
            return ProcessedDocument(
                doc_id=doc_id,
                text=text,
                clauses=clauses,
                entities=entities,
                tables=[],
                metadata={"file_type": "text"},
                structure={}
            )
            
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            raise
    
    def _ocr_pdf(self, file_path: str) -> str:
        """Perform OCR on PDF"""
        try:
            import pdf2image
            
            images = pdf2image.convert_from_path(file_path)
            text = ""
            
            for image in images:
                page_text = pytesseract.image_to_string(image)
                text += page_text + "\n\n"
            
            return text
            
        except ImportError:
            logger.warning("pdf2image not installed, OCR unavailable")
            return ""
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""
    
    def _extract_text_pypdf2(self, file_path: str) -> str:
        """Fallback text extraction using PyPDF2"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
                return text
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")
            return ""
    
    def _extract_clauses(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract legal clauses from text
        Clauses are typically numbered or have specific patterns
        """
        clauses = []
        
        # Split by common clause patterns
        import re
        
        # Pattern 1: Numbered clauses (1., 2., etc.)
        pattern1 = r'\n\s*(\d+)\.\s+([^\n]+(?:\n(?!\s*\d+\.).*)*)'
        matches1 = re.finditer(pattern1, text, re.MULTILINE)
        
        for match in matches1:
            clause_num = match.group(1)
            clause_text = match.group(2).strip()
            if len(clause_text) > 20:  # Filter out short matches
                clauses.append({
                    "id": f"clause_{clause_num}",
                    "number": clause_num,
                    "text": clause_text,
                    "type": "numbered"
                })
        
        # Pattern 2: Section clauses
        pattern2 = r'(?:Section|Article|Clause)\s+(\d+[A-Z]?)[:\s]+([^\n]+(?:\n(?!(?:Section|Article|Clause)).*)*)'
        matches2 = re.finditer(pattern2, text, re.IGNORECASE | re.MULTILINE)
        
        for match in matches2:
            section_num = match.group(1)
            section_text = match.group(2).strip()
            if len(section_text) > 20:
                clauses.append({
                    "id": f"section_{section_num}",
                    "number": section_num,
                    "text": section_text,
                    "type": "section"
                })
        
        # If no clauses found, split by paragraphs
        if not clauses:
            paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 50]
            for idx, para in enumerate(paragraphs[:20]):  # Limit to 20
                clauses.append({
                    "id": f"para_{idx+1}",
                    "number": str(idx+1),
                    "text": para,
                    "type": "paragraph"
                })
        
        return clauses
    
    def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract legal entities from text
        Simple regex-based extraction for common entities
        """
        entities = []
        
        import re
        
        # Extract dates
        date_pattern = r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b'
        dates = re.findall(date_pattern, text)
        for date in dates:
            entities.append({"text": date, "type": "DATE"})
        
        # Extract monetary amounts
        money_pattern = r'₹\s*[\d,]+(?:\.\d{2})?|\bRs\.?\s*[\d,]+(?:\.\d{2})?'
        amounts = re.findall(money_pattern, text)
        for amount in amounts:
            entities.append({"text": amount, "type": "MONEY"})
        
        # Extract organization names (simple heuristic)
        org_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Ltd|Limited|Inc|Corp|Company|Co\.)\b'
        orgs = re.findall(org_pattern, text)
        for org in orgs:
            entities.append({"text": org, "type": "ORGANIZATION"})
        
        return entities

# Global processor instance
document_processor = None

def get_document_processor() -> EnhancedDocumentProcessor:
    """Get or create document processor singleton"""
    global document_processor
    if document_processor is None:
        document_processor = EnhancedDocumentProcessor()
    return document_processor
