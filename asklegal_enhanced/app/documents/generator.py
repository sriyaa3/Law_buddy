from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class LegalDocumentGenerator:
    """Legal document generator for creating contracts, notices, and other legal documents"""
    
    def __init__(self):
        # Document templates
        self.templates = {
            "nda": self._create_nda_template,
            "employment_contract": self._create_employment_contract_template,
            "service_agreement": self._create_service_agreement_template,
            "loan_agreement": self._create_loan_agreement_template,
            "notice": self._create_notice_template
        }
    
    def generate_document(self, template_type: str, details: Dict[str, Any], output_path: str) -> bool:
        """
        Generate legal document from template
        
        Args:
            template_type (str): Type of document template
            details (Dict[str, Any]): Document details
            output_path (str): Output file path
            
        Returns:
            bool: True if successful, False otherwise
        """
        if template_type not in self.templates:
            raise ValueError(f"Unknown template type: {template_type}")
        
        try:
            # Create document using template
            doc = self.templates[template_type](details)
            
            # Save document
            doc.save(output_path)
            print(f"Document saved to: {output_path}")
            return True
            
        except Exception as e:
            print(f"Error generating document: {e}")
            return False
    
    def _create_nda_template(self, details: Dict[str, Any]):
        """Create NDA template"""
        doc = Document()
        
        # Title
        title = doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Effective Date: {details.get('effective_date', datetime.now().strftime('%B %d, %Y'))}")
        doc.add_paragraph()
        
        # Parties
        doc.add_heading('PARTIES', level=1)
        party1 = details.get('disclosing_party', 'Party A')
        party2 = details.get('receiving_party', 'Party B')
        doc.add_paragraph(f"This Non-Disclosure Agreement (the \"Agreement\") is entered into between {party1} (\"Disclosing Party\") and {party2} (\"Receiving Party\").")
        doc.add_paragraph()
        
        # Recitals
        doc.add_heading('RECITALS', level=1)
        doc.add_paragraph("The Disclosing Party possesses certain proprietary and confidential information that it desires to disclose to the Receiving Party for evaluation of a potential business relationship.")
        doc.add_paragraph()
        
        # Definitions
        doc.add_heading('DEFINITIONS', level=1)
        doc.add_paragraph("For purposes of this Agreement, \"Confidential Information\" shall include all information or material that has commercial value and that is not generally known to the public.")
        doc.add_paragraph()
        
        # Obligations
        doc.add_heading('OBLIGATIONS OF RECEIVING PARTY', level=1)
        doc.add_paragraph("The Receiving Party agrees to hold and maintain the Confidential Information in strict confidence and not to disclose it to any third party without the prior written consent of the Disclosing Party.")
        doc.add_paragraph()
        
        # Term
        doc.add_heading('TERM', level=1)
        term = details.get('term', '2 years')
        doc.add_paragraph(f"This Agreement shall remain in effect for a period of {term} from the Effective Date.")
        doc.add_paragraph()
        
        # General Provisions
        doc.add_heading('GENERAL PROVISIONS', level=1)
        doc.add_paragraph("This Agreement shall be binding upon and inure to the benefit of the parties hereto and their respective successors and assigns.")
        doc.add_paragraph()
        
        # Signatures
        doc.add_heading('SIGNATURES', level=1)
        doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
        doc.add_paragraph()
        
        # Signature blocks
        doc.add_table(4, 3)
        table = doc.tables[-1]
        
        # Party 1 signature
        table.cell(0, 0).text = f"{party1}:"
        table.cell(1, 0).text = "Signature: ________________________"
        table.cell(2, 0).text = "Name: ________________________"
        table.cell(3, 0).text = "Date: ________________________"
        
        # Party 2 signature
        table.cell(0, 2).text = f"{party2}:"
        table.cell(1, 2).text = "Signature: ________________________"
        table.cell(2, 2).text = "Name: ________________________"
        table.cell(3, 2).text = "Date: ________________________"
        
        return doc
    
    def _create_employment_contract_template(self, details: Dict[str, Any]):
        """Create employment contract template"""
        doc = Document()
        
        # Title
        title = doc.add_heading('EMPLOYMENT AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Effective Date: {details.get('effective_date', datetime.now().strftime('%B %d, %Y'))}")
        doc.add_paragraph()
        
        # Parties
        doc.add_heading('PARTIES', level=1)
        employer = details.get('employer', 'Employer')
        employee = details.get('employee', 'Employee')
        doc.add_paragraph(f"This Employment Agreement (the \"Agreement\") is entered into between {employer} (\"Employer\") and {employee} (\"Employee\").")
        doc.add_paragraph()
        
        # Position
        doc.add_heading('POSITION AND DUTIES', level=1)
        position = details.get('position', 'Position')
        duties = details.get('duties', 'As assigned by Employer')
        doc.add_paragraph(f"Employee shall serve as {position} and perform such duties as assigned by Employer.")
        doc.add_paragraph()
        
        # Compensation
        doc.add_heading('COMPENSATION', level=1)
        salary = details.get('salary', '$0.00')
        doc.add_paragraph(f"Employee shall receive an annual salary of {salary}, payable in accordance with Employer's standard payroll practices.")
        doc.add_paragraph()
        
        # Benefits
        doc.add_heading('BENEFITS', level=1)
        benefits = details.get('benefits', 'As provided by Employer')
        doc.add_paragraph(f"Employee shall be eligible for benefits including but not limited to: {benefits}")
        doc.add_paragraph()
        
        # Term
        doc.add_heading('TERM', level=1)
        start_date = details.get('start_date', datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph(f"This Agreement shall commence on {start_date} and continue until terminated in accordance with its terms.")
        doc.add_paragraph()
        
        # Termination
        doc.add_heading('TERMINATION', level=1)
        doc.add_paragraph("This Agreement may be terminated by either party with thirty (30) days written notice.")
        doc.add_paragraph()
        
        # Confidentiality
        doc.add_heading('CONFIDENTIALITY', level=1)
        doc.add_paragraph("Employee agrees to maintain the confidentiality of Employer's proprietary information.")
        doc.add_paragraph()
        
        # Signatures
        doc.add_heading('SIGNATURES', level=1)
        doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
        doc.add_paragraph()
        
        # Signature blocks
        doc.add_table(4, 3)
        table = doc.tables[-1]
        
        # Employer signature
        table.cell(0, 0).text = f"{employer}:"
        table.cell(1, 0).text = "Signature: ________________________"
        table.cell(2, 0).text = "Name: ________________________"
        table.cell(3, 0).text = "Title: ________________________"
        
        # Employee signature
        table.cell(0, 2).text = f"{employee}:"
        table.cell(1, 2).text = "Signature: ________________________"
        table.cell(2, 2).text = "Name: ________________________"
        table.cell(3, 2).text = "Date: ________________________"
        
        return doc
    
    def _create_service_agreement_template(self, details: Dict[str, Any]):
        """Create service agreement template"""
        doc = Document()
        
        # Title
        title = doc.add_heading('SERVICE AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Effective Date: {details.get('effective_date', datetime.now().strftime('%B %d, %Y'))}")
        doc.add_paragraph()
        
        # Parties
        doc.add_heading('PARTIES', level=1)
        client = details.get('client', 'Client')
        service_provider = details.get('service_provider', 'Service Provider')
        doc.add_paragraph(f"This Service Agreement (the \"Agreement\") is entered into between {client} (\"Client\") and {service_provider} (\"Service Provider\").")
        doc.add_paragraph()
        
        # Services
        doc.add_heading('SERVICES', level=1)
        services = details.get('services', 'As described in Statement of Work')
        doc.add_paragraph(f"Service Provider shall perform the following services: {services}")
        doc.add_paragraph()
        
        # Term
        doc.add_heading('TERM', level=1)
        term = details.get('term', '1 year')
        doc.add_paragraph(f"This Agreement shall remain in effect for {term} from the Effective Date.")
        doc.add_paragraph()
        
        # Payment
        doc.add_heading('PAYMENT TERMS', level=1)
        payment_terms = details.get('payment_terms', 'Net 30 days')
        doc.add_paragraph(f"Client shall pay Service Provider according to the following terms: {payment_terms}")
        doc.add_paragraph()
        
        # Signatures
        doc.add_heading('SIGNATURES', level=1)
        doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
        doc.add_paragraph()
        
        # Signature blocks
        doc.add_table(4, 3)
        table = doc.tables[-1]
        
        # Client signature
        table.cell(0, 0).text = f"{client}:"
        table.cell(1, 0).text = "Signature: ________________________"
        table.cell(2, 0).text = "Name: ________________________"
        table.cell(3, 0).text = "Title: ________________________"
        
        # Service Provider signature
        table.cell(0, 2).text = f"{service_provider}:"
        table.cell(1, 2).text = "Signature: ________________________"
        table.cell(2, 2).text = "Name: ________________________"
        table.cell(3, 2).text = "Date: ________________________"
        
        return doc
    
    def _create_loan_agreement_template(self, details: Dict[str, Any]):
        """Create loan agreement template"""
        doc = Document()
        
        # Title
        title = doc.add_heading('LOAN AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Effective Date: {details.get('effective_date', datetime.now().strftime('%B %d, %Y'))}")
        doc.add_paragraph()
        
        # Parties
        doc.add_heading('PARTIES', level=1)
        lender = details.get('lender', 'Lender')
        borrower = details.get('borrower', 'Borrower')
        doc.add_paragraph(f"This Loan Agreement (the \"Agreement\") is entered into between {lender} (\"Lender\") and {borrower} (\"Borrower\").")
        doc.add_paragraph()
        
        # Loan Amount
        doc.add_heading('LOAN AMOUNT', level=1)
        loan_amount = details.get('loan_amount', '$0.00')
        doc.add_paragraph(f"Lender agrees to loan Borrower the principal amount of {loan_amount}.")
        doc.add_paragraph()
        
        # Interest Rate
        doc.add_heading('INTEREST RATE', level=1)
        interest_rate = details.get('interest_rate', '0%')
        doc.add_paragraph(f"The loan shall bear interest at the rate of {interest_rate} per annum.")
        doc.add_paragraph()
        
        # Repayment Terms
        doc.add_heading('REPAYMENT TERMS', level=1)
        repayment_terms = details.get('repayment_terms', 'As agreed')
        doc.add_paragraph(f"Borrower shall repay the loan according to the following terms: {repayment_terms}")
        doc.add_paragraph()
        
        # Default
        doc.add_heading('DEFAULT', level=1)
        doc.add_paragraph("If Borrower fails to make any payment when due, Lender may declare the entire unpaid balance immediately due and payable.")
        doc.add_paragraph()
        
        # Signatures
        doc.add_heading('SIGNATURES', level=1)
        doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
        doc.add_paragraph()
        
        # Signature blocks
        doc.add_table(4, 3)
        table = doc.tables[-1]
        
        # Lender signature
        table.cell(0, 0).text = f"{lender}:"
        table.cell(1, 0).text = "Signature: ________________________"
        table.cell(2, 0).text = "Name: ________________________"
        table.cell(3, 0).text = "Date: ________________________"
        
        # Borrower signature
        table.cell(0, 2).text = f"{borrower}:"
        table.cell(1, 2).text = "Signature: ________________________"
        table.cell(2, 2).text = "Name: ________________________"
        table.cell(3, 2).text = "Date: ________________________"
        
        return doc
    
    def _create_notice_template(self, details: Dict[str, Any]):
        """Create notice template"""
        doc = Document()
        
        # Title
        title = doc.add_heading('NOTICE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        notice_date = details.get('notice_date', datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph(f"Date: {notice_date}")
        doc.add_paragraph()
        
        # To
        to_party = details.get('to_party', 'Recipient')
        doc.add_paragraph(f"To: {to_party}")
        doc.add_paragraph()
        
        # From
        from_party = details.get('from_party', 'Sender')
        doc.add_paragraph(f"From: {from_party}")
        doc.add_paragraph()
        
        # Subject
        subject = details.get('subject', 'Notice')
        doc.add_heading(f"Re: {subject}", level=1)
        doc.add_paragraph()
        
        # Notice Content
        notice_content = details.get('notice_content', 'Notice content here')
        doc.add_paragraph(notice_content)
        doc.add_paragraph()
        
        # Required Action
        required_action = details.get('required_action', 'None specified')
        doc.add_paragraph(f"Required Action: {required_action}")
        doc.add_paragraph()
        
        # Response Deadline
        response_deadline = details.get('response_deadline', 'None specified')
        doc.add_paragraph(f"Response Deadline: {response_deadline}")
        doc.add_paragraph()
        
        # Contact Information
        contact_info = details.get('contact_info', 'Contact information here')
        doc.add_paragraph(f"Contact Information: {contact_info}")
        
        return doc

# Global instance
document_generator = LegalDocumentGenerator()